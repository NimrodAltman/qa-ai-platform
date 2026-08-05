"""Generate a richer, fully fictional demo specification (sample_spec.docx).

Run:  python examples/make_sample_spec.py
A realistic-looking process spec — invented data, no real organization.
"""

from pathlib import Path

import docx


def build() -> Path:
    d = docx.Document()
    d.add_heading("אפיון תהליך — אישור בקשת החזר כספי", level=1)
    d.add_paragraph("תיוג משימה: 40100")
    d.add_paragraph(
        "מסמך זה מגדיר את תהליך אישור בקשת החזר כספי ללקוח, כולל תנאי הזכאות, "
        "עדכוני השדות והחריגים."
    )

    d.add_heading("ישויות ושדות", level=2)

    d.add_paragraph('ישות "בקשת החזר" (req_refund):')
    t1 = d.add_table(rows=1, cols=3)
    t1.rows[0].cells[0].text = "שם עברי"
    t1.rows[0].cells[1].text = "שם טכני"
    t1.rows[0].cells[2].text = "הערות"
    for he, tech, note in [
        ("סטטוס בקשה", "req_status", "קוד: 1=חדשה, 2=בטיפול, 3=אושרה, 4=נדחתה"),
        ("סכום מבוקש", "req_amount", "מספרי, חייב להיות גדול מ-0"),
        ("מזהה לקוח", "req_customerid", "מפתח לישות לקוח (crm_customer)"),
        ("תאריך אישור", "req_approveddate", "מתעדכן רק במעבר לסטטוס אושרה (3)"),
        ("מאשר", "req_approverid", "חובה כאשר הבקשה אושרה"),
    ]:
        c = t1.add_row().cells
        c[0].text, c[1].text, c[2].text = he, tech, note

    d.add_paragraph('ישות "לקוח" (crm_customer):')
    t2 = d.add_table(rows=1, cols=3)
    t2.rows[0].cells[0].text = "שם עברי"
    t2.rows[0].cells[1].text = "שם טכני"
    t2.rows[0].cells[2].text = "הערות"
    for he, tech, note in [
        ("סטטוס לקוח", "cust_status", "קוד: 1=פעיל, 2=חסום"),
        ("תקרת החזר", "cust_refundlimit", "מספרי — הסכום המרבי המותר להחזר"),
    ]:
        c = t2.add_row().cells
        c[0].text, c[1].text, c[2].text = he, tech, note

    d.add_heading("חוקים עסקיים", level=2)
    for rule in [
        "בקשה זכאית לאישור רק אם: סטטוס בקשה = 1 (חדשה) או 2 (בטיפול), הלקוח פעיל "
        "(cust_status = 1), והסכום המבוקש (req_amount) קטן או שווה לתקרת ההחזר של הלקוח "
        "(cust_refundlimit).",
        "באישור: req_status עובר ל-3 (אושרה), req_approveddate מתעדכן לתאריך האישור, "
        "ו-req_approverid מאוכלס במזהה המאשר.",
        "אם הסכום המבוקש עולה על תקרת ההחזר — הבקשה נדחית אוטומטית (req_status = 4) "
        "ו-req_approveddate נשאר NULL.",
        "לקוח חסום (cust_status = 2): כל בקשה נדחית (req_status = 4), ללא תלות בסכום.",
        "לא ניתן לאשר בקשה שכבר אושרה (3) או נדחתה (4) — פעולה חוזרת אינה משנה את הרשומה.",
    ]:
        d.add_paragraph(rule, style="List Bullet")

    d.add_heading("ממשקים", level=2)
    d.add_paragraph(
        "עם אישור הבקשה נשלח חיווי למערכת התשלומים. שליחת החיווי עצמה מחוץ להיקף בדיקה זו, "
        "אך יש לוודא שנרשמת שורת לוג על שליחת החיווי."
    )

    out = Path(__file__).with_name("sample_spec.docx")
    d.save(out)
    return out


if __name__ == "__main__":
    build()
    print("Wrote sample_spec.docx")
