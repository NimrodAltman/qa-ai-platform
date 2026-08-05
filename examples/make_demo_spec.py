"""Generate a fully fictional demo specification (demo_spec.docx).

Run:  python examples/make_demo_spec.py
The content is invented for demonstration — no real organizational data.
"""

from pathlib import Path

import docx


def build() -> Path:
    document = docx.Document()
    document.add_heading("אפיון: מנגנון עדכון סטטוס הזמנה", level=1)
    document.add_paragraph(
        "תיוג 40012: כאשר מתקבל קובץ אישור מספק, יש לעדכן את סטטוס ההזמנה."
    )
    document.add_paragraph(
        'ישות "הזמנה" (demo_order) מכילה את השדות הבאים:'
    )

    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "שם עברי"
    table.rows[0].cells[1].text = "שם טכני"
    table.rows[0].cells[2].text = "הערות"
    rows = [
        ("סטטוס הזמנה", "demo_orderstatus", "ערכי קוד: ממתין (1), אושר (2), נדחה (3)"),
        ("מזהה ספק", "demo_supplierid", "מפתח לישות ספק (demo_supplier)"),
        ("תאריך אישור", "demo_approveddate", "מתעדכן רק במעבר לסטטוס אושר"),
    ]
    for he, tech, note in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = he, tech, note

    document.add_paragraph(
        "חוק עסקי: רשומה נכללת בעדכון רק אם demo_orderstatus = 1 (ממתין) "
        "וקיים demo_supplierid. אם הקובץ תקין — הסטטוס עובר ל-2 (אושר) "
        "ו-demo_approveddate מתעדכן לתאריך הקליטה."
    )

    out = Path(__file__).with_name("demo_spec.docx")
    document.save(out)
    return out


if __name__ == "__main__":
    build()
    print("Wrote demo_spec.docx")
