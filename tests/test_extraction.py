"""Tests for document extraction using synthetic .docx and .xlsx files."""

import docx
import pytest
from openpyxl import Workbook

from qa_agents.extraction import extract


def test_extract_docx_keeps_paragraphs_and_tables(tmp_path):
    document = docx.Document()
    document.add_heading("אפיון תהליך", level=1)
    document.add_paragraph("תיוג 12345: עדכון שדה סטטוס.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "שדה"
    table.rows[0].cells[1].text = "סכמה"
    table.rows[1].cells[0].text = "סטטוס"
    table.rows[1].cells[1].text = "new_status"
    path = tmp_path / "spec.docx"
    document.save(path)

    text = extract(path)

    assert "אפיון תהליך" in text
    assert "תיוג 12345" in text
    assert "שדה | סכמה" in text
    assert "סטטוס | new_status" in text


def test_extract_xlsx_lists_sheets_and_rows(tmp_path):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "אפיון"
    sheet.append(["תיוג", "שדה", "ערך"])
    sheet.append([12345, "new_status", 100000000])
    sheet.append([None, None, None])  # empty row is skipped
    path = tmp_path / "spec.xlsx"
    workbook.save(path)

    text = extract(path)

    assert "# Sheet: אפיון" in text
    assert "תיוג | שדה | ערך" in text
    assert "12345 | new_status | 100000000" in text


def test_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        extract(tmp_path / "nope.docx")


def test_unsupported_extension_raises(tmp_path):
    path = tmp_path / "spec.txt"
    path.write_text("hello", encoding="utf-8")
    with pytest.raises(ValueError):
        extract(path)
